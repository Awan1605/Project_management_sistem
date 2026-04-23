"""
Document Content Extraction Utility
=====================================
Extract text content from various file formats (PDF, DOCX, XLSX, TXT)
for use in RAG knowledge base.
"""

import os
import logging
from django.conf import settings

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path):
    """
    Extract text content from a file based on its extension.
    
    Supported formats:
    - PDF (.pdf) - using PyPDF2
    - Word (.docx) - using python-docx
    - Excel (.xlsx, .xls) - using openpyxl
    - Text (.txt, .md, .csv) - plain text
    
    Args:
        file_path: Absolute path to the file
        
    Returns:
        str: Extracted text content, or empty string if extraction fails
    """
    if not os.path.exists(file_path):
        logger.warning(f"[Doc Extract] File not found: {file_path}")
        return ""
    
    # Get file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    try:
        if ext == '.pdf':
            return _extract_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return _extract_docx(file_path)
        elif ext in ['.xlsx', '.xls']:
            return _extract_excel(file_path)
        elif ext in ['.txt', '.md', '.csv', '.log', '.json', '.xml', '.html']:
            return _extract_text(file_path)
        else:
            logger.info(f"[Doc Extract] Unsupported format: {ext}")
            return ""
            
    except Exception as e:
        logger.error(f"[Doc Extract] Error extracting {file_path}: {e}")
        return ""


def _extract_pdf(file_path):
    """Extract text from PDF file using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[Page {page_num}]\n{page_text}")
        
        return "\n\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"[PDF Extract] Error: {e}")
        return ""


def _extract_docx(file_path):
    """Extract text from Word document using python-docx."""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"[DOCX Extract] Error: {e}")
        return ""


def _extract_excel(file_path):
    """Extract text from Excel file using openpyxl."""
    try:
        from openpyxl import load_workbook
        
        wb = load_workbook(file_path, data_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")
            
            for row in ws.iter_rows(values_only=True):
                # Filter out empty rows
                non_empty = [str(cell) for cell in row if cell is not None]
                if non_empty:
                    text_parts.append(" | ".join(non_empty))
        
        return "\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"[Excel Extract] Error: {e}")
        return ""


def _extract_text(file_path):
    """Extract text from plain text file."""
    try:
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    # Limit to first 10000 chars to avoid huge documents
                    if len(content) > 10000:
                        content = content[:10000] + "\n\n[... truncated ...]"
                    return content
            except UnicodeDecodeError:
                continue
        
        logger.warning(f"[Text Extract] Could not decode file with any encoding")
        return ""
        
    except Exception as e:
        logger.error(f"[Text Extract] Error: {e}")
        return ""


def get_file_summary(file_path, max_length=2000):
    """
    Get a summarized version of document content for RAG.
    Truncates to max_length while keeping important parts.
    
    Args:
        file_path: Path to the file
        max_length: Maximum length of summary
        
    Returns:
        str: Summarized document content
    """
    full_text = extract_text_from_file(file_path)
    
    if not full_text:
        return "[No text content could be extracted]"
    
    # If text is short enough, return as-is
    if len(full_text) <= max_length:
        return full_text
    
    # Otherwise, truncate intelligently
    # Keep the first part (usually contains title/intro)
    summary = full_text[:max_length]
    
    # Try to cut at a sentence boundary
    last_period = summary.rfind('.')
    if last_period > max_length * 0.7:  # If we can find a period in the last 30%
        summary = summary[:last_period + 1]
    
    summary += "\n\n[... document truncated ...]"
    
    return summary
